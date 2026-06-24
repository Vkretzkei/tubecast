#!/usr/bin/env python3
"""
TubeCast
Porta local: 7474  |  Cloud: usa variável PORT
"""

import io
import json
import mimetypes
import os
import platform
import re
import subprocess
import sys
import tempfile
import time
import uuid
import webbrowser
from http.server import BaseHTTPRequestHandler, HTTPServer
from pathlib import Path
from socketserver import ThreadingMixIn
from threading import Lock, Timer

# ── Configuração ──────────────────────────────────────────────────────────────
PORT       = int(os.environ.get("PORT", 7474))
CLOUD_MODE = "PORT" in os.environ          # Render/Railway define PORT automaticamente
STATIC_DIR = Path(__file__).parent
IS_WINDOWS = platform.system() == "Windows"

ASSEMBLYAI_NATIVE   = {".mp3", ".mp4", ".m4a", ".wav", ".ogg", ".flac", ".webm"}
ACCEPTED_EXTENSIONS = ASSEMBLYAI_NATIVE | {".wma", ".mkv", ".avi", ".mov", ".opus", ".weba", ".aac"}

RE_PCT   = re.compile(r"(\d+\.?\d*)%")
RE_SPEED = re.compile(r"at\s+([\d.]+\s*\S+/s)")
RE_SAFE  = re.compile(r'[<>:"/\\|?*]')

# ── Armazenamento temporário de resultados (max 10 min) ───────────────────────
_results: dict[str, dict] = {}
_results_lock = Lock()

def store_result(data: bytes, filename: str) -> str:
    rid = str(uuid.uuid4())
    with _results_lock:
        _results[rid] = {"data": data, "filename": filename, "ts": time.time()}
        # limpa resultados com mais de 10 minutos
        stale = [k for k, v in _results.items() if time.time() - v["ts"] > 600]
        for k in stale:
            del _results[k]
    return rid

def pop_result(rid: str) -> dict | None:
    with _results_lock:
        return _results.pop(rid, None)


# ── Servidor multi-thread ─────────────────────────────────────────────────────
class ThreadingHTTPServer(ThreadingMixIn, HTTPServer):
    daemon_threads = True


# ── Helpers ───────────────────────────────────────────────────────────────────
def sse(data: dict) -> bytes:
    return ("data:" + json.dumps(data, ensure_ascii=False) + "\n\n").encode("utf-8")

def ensure_deps():
    for pkg, pip_name in [("assemblyai", "assemblyai"), ("docx", "python-docx")]:
        try:
            __import__(pkg)
        except ImportError:
            print(f"Instalando {pip_name}...")
            subprocess.run([sys.executable, "-m", "pip", "install", pip_name, "-q"], check=True)

def check_tool(name: str) -> bool:
    cmd = ["where", name] if IS_WINDOWS else ["which", name]
    return subprocess.run(cmd, capture_output=True).returncode == 0

def convert_to_mp3(input_path: Path, output_path: Path) -> bool:
    result = subprocess.run([
        "ffmpeg", "-y", "-i", str(input_path),
        "-vn", "-ar", "16000", "-ac", "1", "-b:a", "48k",
        str(output_path)
    ], capture_output=True)
    return result.returncode == 0

def stream_download(url: str, output_template: str, emit_fn,
                    quality_kbps: str = "320", audio_format: str = "mp3") -> bool:
    cmd = [
        "yt-dlp", "--extract-audio",
        "--audio-format", audio_format,
        "--audio-quality", quality_kbps + "K",
        "--newline", "--no-playlist",
        "--concurrent-fragments", "4",
        "--no-warnings", "--no-part",
        "--output", output_template,
        url,
    ]
    try:
        proc = subprocess.Popen(
            cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
            text=True, bufsize=1, encoding="utf-8", errors="replace"
        )
        for line in proc.stdout:
            line = line.strip()
            if not line:
                continue
            if "[download]" in line:
                pct_m = RE_PCT.search(line)
                spd_m = RE_SPEED.search(line)
                if pct_m:
                    emit_fn({"type": "progress",
                             "pct": round(float(pct_m.group(1)), 1),
                             "speed": spd_m.group(1) if spd_m else ""})
            elif "[ExtractAudio]" in line or "[ffmpeg]" in line:
                emit_fn({"type": "converting"})
        proc.wait()
        return proc.returncode == 0
    except FileNotFoundError:
        emit_fn({"type": "error", "msg": "yt-dlp não encontrado. Instala com: pip install yt-dlp"})
        return False
    except Exception as e:
        emit_fn({"type": "error", "msg": str(e)})
        return False

def build_docx(title: str, source_label: str, lang: str, transcript_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_par.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    doc.add_paragraph()

    lang_labels = {"pt": "Português", "en": "Inglês", "es": "Espanhol", "": "Automático"}
    for meta_text in [
        f"Fonte: {source_label}",
        f"Idioma: {lang_labels.get(lang, lang)}  |  Modelo: AssemblyAI (Best)"
    ]:
        m  = doc.add_paragraph()
        mr = m.add_run(meta_text)
        mr.font.size = Pt(9)
        mr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    sep = doc.add_paragraph()
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    doc.add_paragraph()

    sentences = re.split(r'(?<=[.!?])\s+', transcript_text.strip())
    chunk = []
    for sentence in sentences:
        chunk.append(sentence)
        if len(chunk) >= 4:
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = par.add_run(" ".join(chunk))
            r.font.size = Pt(12)
            chunk = []
    if chunk:
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = par.add_run(" ".join(chunk))
        r.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def transcribe_audio(audio_path: Path, api_key: str, lang: str, emit_fn) -> str | None:
    emit_fn({"type": "transcribing"})
    try:
        import assemblyai as aai
        aai.settings.api_key = api_key
        config_kwargs = {"speech_model": aai.SpeechModel.best}
        if lang in ("pt", "en", "es"):
            config_kwargs["language_code"] = lang
        else:
            config_kwargs["language_detection"] = True
        transcript = aai.Transcriber(config=aai.TranscriptionConfig(**config_kwargs)).transcribe(str(audio_path))
        if transcript.status == aai.TranscriptStatus.error:
            emit_fn({"type": "error", "msg": f"Erro AssemblyAI: {transcript.error}"})
            return None
        return transcript.text or ""
    except Exception as e:
        emit_fn({"type": "error", "msg": f"Erro na transcrição: {str(e)}"})
        return None


# ── Handler ───────────────────────────────────────────────────────────────────
class Handler(BaseHTTPRequestHandler):

    def log_message(self, format, *args):
        pass

    def send_cors(self):
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers",
                         "Content-Type, X-Api-Key, X-Lang")

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_cors()
        self.end_headers()

    def do_GET(self):
        path = self.path.split("?")[0]

        if path == "/config":
            self._json({
                "defaultDir": str(DEFAULT_OUTPUT_DIR),
                "apiKey": os.environ.get("ASSEMBLYAI_API_KEY", ""),
            })
            return

        if path.startswith("/result/"):
            self._serve_result(path[8:])
            return

        if path == "/" or path == "/index.html":
            file_path = STATIC_DIR / "index.html"
        else:
            file_path = STATIC_DIR / path.lstrip("/")

        if file_path.exists() and file_path.is_file():
            mime, _ = mimetypes.guess_type(str(file_path))
            content = file_path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", mime or "application/octet-stream")
            self.send_header("Content-Length", str(len(content)))
            self.end_headers()
            self.wfile.write(content)
        else:
            self.send_response(404)
            self.end_headers()

    def do_POST(self):
        routes = {
            "/download":        self.handle_download,
            "/transcribe":      self.handle_transcribe,
            "/transcribe-file": self.handle_transcribe_file,
        }
        handler = routes.get(self.path)
        if handler:
            handler()
        else:
            self.send_response(404)
            self.end_headers()

    # ── GET /result/<id> ─────────────────────────────────────────────────── #
    def _serve_result(self, rid: str):
        result = pop_result(rid)
        if not result:
            self.send_response(404)
            self.end_headers()
            return
        mime, _ = mimetypes.guess_type(result["filename"])
        self.send_response(200)
        self.send_cors()
        self.send_header("Content-Type", mime or "application/octet-stream")
        self.send_header("Content-Disposition",
                         f'attachment; filename="{result["filename"]}"')
        self.send_header("Content-Length", str(len(result["data"])))
        self.end_headers()
        self.wfile.write(result["data"])

    # ── /download ─────────────────────────────────────────────────────────── #
    def handle_download(self):
        body    = self._read_json()
        url     = body.get("url", "").strip()
        quality = str(body.get("quality", "320"))

        self._start_sse()

        with tempfile.TemporaryDirectory() as tmpdir:
            ok = stream_download(url, str(Path(tmpdir) / "%(title)s.%(ext)s"),
                                 self._emit, quality_kbps=quality)
            if not ok:
                self._emit({"type": "error", "msg": "Erro ao baixar. Verifique o link."})
                return

            files = list(Path(tmpdir).glob("*.mp3")) or list(Path(tmpdir).glob("*.*"))
            if not files:
                self._emit({"type": "error", "msg": "Arquivo não encontrado após download."})
                return

            audio_path = files[0]
            filename   = audio_path.name
            data       = audio_path.read_bytes()

        rid = store_result(data, filename)
        self._emit({"type": "done", "id": rid, "filename": filename})

    # ── /transcribe ───────────────────────────────────────────────────────── #
    def handle_transcribe(self):
        body    = self._read_json()
        url     = body.get("url", "").strip()
        api_key = body.get("apiKey", "").strip()
        lang    = body.get("lang", "pt").strip()

        self._start_sse()

        with tempfile.TemporaryDirectory() as tmpdir:
            ok = stream_download(url, str(Path(tmpdir) / "%(title)s.%(ext)s"),
                                 self._emit, quality_kbps="48", audio_format="mp3")
            if not ok:
                return

            files = list(Path(tmpdir).glob("*.mp3")) or list(Path(tmpdir).glob("*.*"))
            if not files:
                self._emit({"type": "error", "msg": "Arquivo de áudio não encontrado."})
                return

            audio_path = files[0]
            title      = audio_path.stem
            text       = transcribe_audio(audio_path, api_key, lang, self._emit)

        if text is None:
            return

        try:
            docx_bytes = build_docx(title, url, lang, text)
        except Exception as e:
            self._emit({"type": "error", "msg": f"Erro ao criar Word: {e}"})
            return

        safe_title = RE_SAFE.sub('_', title)[:80]
        filename   = f"{safe_title}.docx"
        rid = store_result(docx_bytes, filename)
        self._emit({"type": "done", "id": rid, "filename": filename})

    # ── /transcribe-file ──────────────────────────────────────────────────── #
    def handle_transcribe_file(self):
        api_key      = self.headers.get("X-Api-Key", "").strip()
        lang         = self.headers.get("X-Lang", "pt").strip()
        content_type = self.headers.get("Content-Type", "")
        length       = int(self.headers.get("Content-Length", 0))

        self._start_sse()
        self._emit({"type": "uploading"})

        try:
            import email
            raw = self.rfile.read(length)
            msg = email.message_from_bytes(
                f"Content-Type: {content_type}\r\n\r\n".encode() + raw
            )
            file_data, filename = None, "arquivo"
            for part in msg.walk():
                if part.get_filename():
                    filename  = part.get_filename()
                    file_data = part.get_payload(decode=True)
                    break

            if not file_data:
                self._emit({"type": "error", "msg": "Arquivo não recebido."})
                return

            file_ext = Path(filename).suffix.lower()
            if file_ext not in ACCEPTED_EXTENSIONS:
                self._emit({"type": "error", "msg": f"Formato não suportado: {file_ext}"})
                return
        except Exception as e:
            self._emit({"type": "error", "msg": f"Erro ao ler arquivo: {e}"})
            return

        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = Path(tmpdir) / filename
            input_path.write_bytes(file_data)

            if file_ext in ASSEMBLYAI_NATIVE:
                audio_path = input_path
            else:
                self._emit({"type": "converting"})
                mp3_path = Path(tmpdir) / f"{Path(filename).stem}.mp3"
                if not convert_to_mp3(input_path, mp3_path):
                    self._emit({"type": "error", "msg": "Falha ao converter o arquivo."})
                    return
                audio_path = mp3_path

            title = Path(filename).stem
            text  = transcribe_audio(audio_path, api_key, lang, self._emit)

        if text is None:
            return

        try:
            docx_bytes = build_docx(title, f"Arquivo: {filename}", lang, text)
        except Exception as e:
            self._emit({"type": "error", "msg": f"Erro ao criar Word: {e}"})
            return

        safe_title = RE_SAFE.sub('_', title)[:80]
        filename_out = f"{safe_title}.docx"
        rid = store_result(docx_bytes, filename_out)
        self._emit({"type": "done", "id": rid, "filename": filename_out})

    # ── Helpers ───────────────────────────────────────────────────────────── #
    def _read_json(self) -> dict:
        length = int(self.headers.get("Content-Length", 0))
        return json.loads(self.rfile.read(length))

    def _start_sse(self):
        self.send_response(200)
        self.send_cors()
        self.send_header("Content-Type", "text/event-stream")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("X-Accel-Buffering", "no")
        self.end_headers()

    def _emit(self, data: dict):
        try:
            self.wfile.write(sse(data))
            self.wfile.flush()
        except Exception:
            pass


# ── Startup ───────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if not CLOUD_MODE:
        print("Verificando dependências...")
        ensure_deps()

        missing = [dep for dep in ["yt-dlp", "ffmpeg"] if not check_tool(dep)]
        if missing:
            print(f"\nAviso: {', '.join(missing)} não encontrado(s) no PATH.")
            print("  yt-dlp:  pip install yt-dlp")
            print("  ffmpeg:  https://ffmpeg.org/download.html\n")

    host = "0.0.0.0" if CLOUD_MODE else "localhost"
    url  = f"http://localhost:{PORT}"

    print(f"TubeCast rodando em {url}")
    if not CLOUD_MODE:
        print("Ctrl+C para parar.\n")
        Timer(0.5, lambda: webbrowser.open(url)).start()

    server = ThreadingHTTPServer((host, PORT), Handler)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nServidor encerrado.")
